from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from openpyxl import load_workbook

document = Document()  # Crée un nouveau document Word

def create_element(name):
    return OxmlElement(name)  # Crée un élément XML avec le nom spécifié
        
def create_attribute(element, name, value):
    element.set(ns.qn(name), value)  # Attribue une valeur à un attribut d'un élément XML
        
def add_page_number(run):
    fldStart = create_element('w:fldChar')  # Commence un champ dans le document
    create_attribute(fldStart, 'w:fldCharType', 'begin')  # Spécifie que c'est le début du champ

    instrText = create_element('w:instrText')  # Crée un texte d'instruction pour le champ
    create_attribute(instrText, 'xml:space', 'preserve')  # Préserve les espaces dans le texte
    instrText.text = "PAGE"  # Définit le texte d'instruction à "PAGE" pour le numéro de page

    fldChar1 = create_element('w:fldChar')  # Crée un caractère de champ
    create_attribute(fldChar1, 'w:fldCharType', 'separate')  # Sépare le champ du texte

    fldChar2 = create_element('w:t')  # Crée un texte pour le champ
    fldChar2.text = "2"  # Met le texte du champ à "2"

    fldEnd = create_element('w:fldChar')  # Termine le champ
    create_attribute(fldEnd, 'w:fldCharType', 'end')  # Marque la fin du champ

    # Ajoute tous les éléments créés au run passé en argument
    run._r.append(fldStart)
    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)
    run._r.append(fldEnd)

def Create_footer_header():
    date = datetime.today().strftime('%Y-%m-%d')  # Obtient la date actuelle formatée

    section = document.sections[0]  # Accède à la première section du document
    header = section.header  # Accède à l'entête de la section
    footer = section.footer  # Accède au pied de page de la section
    largeur_tableau = section.page_width - section.left_margin - section.right_margin  # Calcule la largeur disponible pour le tableau
    htable = header.add_table(rows=2, cols=3, width=largeur_tableau)  # Ajoute un tableau à l'entête
    htable.style = 'Table Grid'  # Applique un style de grille au tableau
    # Remplit la première ligne du tableau avec des titres
    hrow = htable.rows[0].cells
    hrow[0].text = 'Auteur'
    hrow[1].text = 'Titre'
    hrow[2].text = 'Date'
    # Remplit la deuxième ligne du tableau avec des valeurs
    hrow = htable.rows[1].cells
    hrow[0].text = 'Automated'
    hrow[1].text = 'Automated Pentest Report'
    hrow[2].text = date

    add_page_number(footer.paragraphs[0].add_run())  # Ajoute un numéro de page au pied de page
    footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Alignement à droite pour le pied de page

def content_table():
    document.add_heading('Table des matières', level=1)  # Ajoute un titre pour la table des matières
    paragraph = document.add_paragraph()  # Ajoute un paragraphe sous le titre
    run = paragraph.add_run()  # Ajoute un run au paragraphe
    fldChar = OxmlElement('w:fldChar')  # Crée un champ pour la table des matières
    fldChar.set(qn('w:fldCharType'), 'begin')  # Commence le champ
    instrText = OxmlElement('w:instrText')  # Texte d'instruction pour le champ
    instrText.set(qn('xml:space'), 'preserve')  # Préserve les espaces
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Définit le texte pour la génération de la TOC

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')  # Sépare le champ du reste
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."  # Indique comment mettre à jour le champ
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')  # Termine le champ

    # Ajoute tous les éléments au run
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

def read_file(path):
    # Lit et retourne le contenu complet d'un fichier texte.
    with open(path, 'r', encoding='utf-8') as fichier:
        contenu = fichier.read()
    return contenu

def line_break(number):
    # Ajoute un nombre spécifié de paragraphes vides (sauts de ligne) au document.
    for _ in range(number):
        document.add_paragraph()

def Get_Table(path):
    # Charge un fichier Excel et insère son contenu sous forme de tableau dans le document Word.
    wb = load_workbook(path)  # Charge le classeur Excel.
    ws = wb.active  # Sélectionne la feuille active.
    table = document.add_table(rows=ws.max_row, cols=ws.max_column)  # Crée un tableau dans le document Word avec le nombre approprié de lignes et de colonnes.
    table.style = 'Table Grid'  # Applique un style de grille au tableau pour une meilleure visibilité.
    for row_idx, row in enumerate(ws.iter_rows(), 1):  # Itère sur chaque ligne du tableau Excel.
        for col_idx, cell in enumerate(row, 1):  # Itère sur chaque cellule de la ligne.
            if cell.value:
                table.cell(row_idx - 1, col_idx - 1).text = str(cell.value)  # Transfère la valeur de la cellule Excel dans la cellule du tableau Word.
            # Les lignes commentées ci-dessous permettent de gérer les cellules vides en ajoutant un espace, mais sont actuellement désactivées.
            '''else:
                table.cell(row_idx - 1, col_idx - 1).text = " "'''

def Get_Text(path):
    # Lit le texte d'un fichier et l'ajoute comme paragraphe dans le document Word.
    text = read_file(path)  # Utilise la fonction read_file pour obtenir le contenu du fichier.
    document.add_paragraph(text)  # Ajoute ce contenu comme un nouveau paragraphe dans le document.

def create_report(host):
    # Fonction principale pour créer le rapport de pentest.
    Create_footer_header()  # Ajoute un en-tête et un pied de page personnalisés au document.

    title = document.add_heading('Pentest automatic report', level=0)  # Ajoute un titre principal au document.
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centre le titre.
    title.paragraph_format.space_before = Pt(250)  # Ajoute un espace avant le titre pour l'espacement.

    document.add_paragraph('Project realised by Guillaume, Yassir and Victor')  # Ajoute une mention des auteurs.

    document.add_page_break()  # Insère un saut de page.

    content_table()  # Ajoute une table des matières.

    document.add_page_break()  # Insère un autre saut de page.

    # Ajoute différents sections du rapport avec leurs titres et contenus associés.
    document.add_heading('Context:', level=1)
    Get_Text('./Doc/Text/context.txt')

    document.add_heading('Introduction:', level=1)
    Get_Text('./Doc/Text/introduction.txt')

    document.add_page_break()

    document.add_heading('Rapport du Pentest: ', level=1)
    Get_Text('./Doc/Text/Report.txt')
    document.add_heading('Rapport des vulnérabilités: ', level=2)
    document.add_picture(f'./Doc/Report/PNG/Vulns/{host}.png')
    document.add_heading('Rapport des exploits: ', level=2)
    document.add_picture(f'./Doc/Report/PNG/Exploit/{host}.png')

    document.add_heading('Reconaissance: ', level=1)
    Get_Text('./Doc/Text/Recon.txt')
    Get_Table(f"./Doc/Report/Excel/Recon/{host}.xlsx")

    document.add_heading('Rapport de vulnérabilité: ', level=1)
    Get_Text('./Doc/Text/CVE.txt')
    document.add_heading('Rapport de vulnérabilité: ', level=2)
    Get_Table(f"./Doc/Report/Excel/Vulns/{host}.xlsx")
    document.add_heading('Rapport de vulnérabilité: ', level=2)
    Get_Table(f"./Doc/Report/Excel/ExploitDB/{host}/exploit.xlsx")

    document.add_heading('Remédiation: ', level=1)
    Get_Text('./Doc/Text/Remed.txt')
    Get_Table(f"./Doc/Report/Excel/Remed/{host}.xlsx")

    document.save(f'./Doc/Report/Docx/Final_Report_{host}.docx')  # Sauvegarde le document final.
