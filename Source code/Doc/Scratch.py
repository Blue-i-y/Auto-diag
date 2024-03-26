from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from openpyxl import load_workbook


def create_element(name):
    return OxmlElement(name)
        
def create_attribute(element, name, value):
    element.set(ns.qn(name), value)
        
def add_page_number(run):
    fldStart = create_element('w:fldChar')
    create_attribute(fldStart, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'separate')

    fldChar2 = create_element('w:t')
    fldChar2.text = "2"

    fldEnd = create_element('w:fldChar')
    create_attribute(fldEnd, 'w:fldCharType', 'end')

    run._r.append(fldStart)

    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)

    run._r.append(fldEnd)

def Create_footer_header():
    global date, document, section
    date = datetime.today().strftime('%Y-%m-%d')

    section = document.sections[0]
    header = section.header
    footer = section.footer
    largeur_tableau = section.page_width - section.left_margin - section.right_margin
    htable = header.add_table(rows=2, cols=3, width=largeur_tableau)
    htable.style = 'Table Grid'
    hrow = htable.rows[0].cells
    hrow[0].text = 'Auteur'
    hrow[1].text = 'Titre'
    hrow[2].text =  'Date'
    hrow = htable.rows[1].cells
    hrow[0].text = 'Automated'
    hrow[1].text = 'Automated Pentest Report'
    hrow[2].text =  date

    add_page_number(footer.paragraphs[0].add_run())
    footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def content_table():
    document.add_heading('Table des matières', level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p

def read_file(path):
    with open(path, 'r', encoding='utf-8') as fichier:
        contenu = fichier.read()
    return contenu

def line_break(number):
    for _ in range(number):
        document.add_paragraph()

def Get_Table(Path):
    wb = load_workbook(Path)
    ws = wb.active
    table = document.add_table(rows=ws.max_row + 1, cols=ws.max_column + 1)
    table.style = 'Table Grid'
    for row_idx, row in enumerate(ws.iter_rows(), 1):
        for col_idx, cell in enumerate(row, 1):
            table.cell(row_idx, col_idx).text = str(cell.value)

def Get_Text(Path):
    text = read_file(Path)
    document.add_paragraph(text)

document = Document()
Create_footer_header()

title = document.add_heading('Pentest automatic report', level=0)
title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.paragraph_format.space_before = Pt(250)
document.add_paragraph('Project realised by Guillaume, Yassir and Victor')

document.add_page_break()

content_table()

document.add_page_break()

document.add_heading('Context:', level=1)
Get_Text('Text/context.txt')

document.add_heading('Introduction:', level=1)
Get_Text('Text/introduction.txt')

document.add_page_break()

document.add_heading('Rapport du Pentest: ', level=1)
Get_Text('Text/Report.txt')
#Report_Graph()
Get_Table('Report/Excel/Rapport.xlsx')

document.add_heading('Reconaissance: ', level=1)
Get_Text('Text/Recon.txt')
Get_Table('Report/Excel/Recon.xlsx')

document.add_heading('Rapport de vulnérabilité: ', level=1)
Get_Text('Text/CVE.txt')
Get_Table('Report/Excel/CVE.xlsx')

document.add_heading('Remédiation: ', level=1)
Get_Text('Text/Remed.txt')
Get_Table("Report/Excel/Remed.xlsx")

document.save('Report/Docx/scratch.docx')
